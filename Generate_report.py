'''
Created on Aug 5, 2014

@author: vvaka
'''
from docx import Document
from docx.shared import Inches
import time

def add_firstpage(document):
    documentHeading = ('80WXC\nOptical Network Health Check Report\nVersion 1.0\n')
    titles = document.add_heading(documentHeading, level = 0)
    titles.alignment = 1

    titlepagePara = ('Corporate Headquarters:\nCisco\n170 West Tasman Drive\nSan Jose, CA 95134-1706\nUSA\nhttp://www.cisco.com\n Tel:    408 526-4000\n800 553-NETS(6387)\nFax:   408 526-4100')

    titlePara = document.add_paragraph()
    titlePara.alignment = 1
    run = titlePara.add_run(titlepagePara)
    run.bold = True
    run.fontsize = '10'
    
    return document


def main():
    
    document = Document()
    docDate = time.strftime("%Y-%m-%d.docx")
    
    document = add_firstpage(document)
    
    
    document.save('80WXC'+docDate)
    
if __name__ == '__main__':
  main()